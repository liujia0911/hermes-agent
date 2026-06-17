"""将陶香居合同 Markdown 生成规范排版的 Word 文档。"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ---- 页面设置 ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
    return h


def add_para(text, bold=False, alignment=None, font_size=11, font_name='宋体'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if alignment is not None:
        p.alignment = alignment
    return p


def set_cell_text(cell, text, bold=False, font_size=10, alignment=None, font_name='宋体'):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if alignment is not None:
        p.alignment = alignment


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, font_size=10,
                      font_name='黑体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            set_cell_text(table.rows[r + 1].cells[c], val, font_size=10)
    if col_widths:
        for row_obj in table.rows:
            for i, w in enumerate(col_widths):
                row_obj.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_line():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('—' * 30)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)


# ============================
# 正文
# ============================

add_heading_styled('陶香居私房菜 · 服务合同', level=1)
add_para('合同编号：XG-2026-0601', bold=True)
add_para('签署日期：2026年____月____日', bold=True)
add_line()

add_heading_styled('甲方（服务方）：昕冠科技', level=2)
add_heading_styled('乙方（客户方）：陶香居私房菜', level=2)
add_line()

# 一、服务内容
add_heading_styled('一、服务内容', level=2)
add_para('甲方为乙方提供以下服务：')

add_heading_styled('1. 抖音获客服务', level=3)
add_table(
    ['服务模块', '说明'],
    [
        ['视频获客', '日更视频制作 + 平台发布'],
        ['广告投流', '广告计划搭建 + 盯盘实时优化'],
        ['SEO 搜索', '抖音搜索引擎排名优化'],
        ['GEO 推荐', '豆包 / DeepSeek / Kimi 等 AI 搜索推荐覆盖'],
        ['来客后台维护', '抖音来客商家后台日常运营维护'],
    ],
    col_widths=[4, 10]
)

add_heading_styled('2. 首期团购达人服务（本合同期内）', level=3)
add_table(
    ['服务项', '数量', '说明'],
    [
        ['达人探店', '8 名', '本地美食 / 生活达人到店拍摄 + 发布'],
        ['矩阵种草', '18 个视频', '多账号矩阵内容分发，覆盖同城流量'],
        ['视频素材包', '100–200 条', '门店环境 + 菜品 + 后厨全套专业拍摄素材'],
        ['定期拍摄新品', '按需', '合同期内新菜品定期上门拍摄'],
    ],
    col_widths=[3.5, 2.5, 8]
)
add_line()

# 二、费用与支付
add_heading_styled('二、费用与支付', level=2)
add_heading_styled('费用明细', level=3)
add_table(
    ['费用项', '金额', '说明'],
    [
        ['月度服务费', '4,000 元/月', '含 AI 获客系统 + 配套内容服务，无额外费用'],
        ['团购搭建服务费', '3,000 元', '抖音 / 美团等平台团购套餐搭建 + 上架（一次性）'],
        ['后续团购服务费扣点', '抖音平台 2.5% + 服务商 8%，共计 10.5%', '按实际核销金额结算'],
    ],
    col_widths=[3.5, 5, 5.5]
)

add_heading_styled('支付方式', level=3)
add_para('签约周期：季度签（3 个月一签）', bold=True)
add_para('首期应付：15,000 元（月费 4,000 × 3 个月 + 团购服务 3,000 元）', bold=True)
add_para('续费：每季度到期前 7 日支付下一季度费用 12,000 元', bold=True)
add_para('支付账户：_______________', bold=True)
add_line()

# 三、服务周期
add_heading_styled('三、服务周期', level=2)
add_table(
    ['节点', '时限'],
    [
        ['合同生效', '首笔款项到账之日起'],
        ['首次交付', '合同生效后 7 个工作日内完成系统部署 + 来客后台接入'],
        ['达人探店', '合同生效后 10 个工作日内完成'],
        ['矩阵种草', '合同期内持续分发，共 18 个视频'],
        ['视频素材', '合同生效后 15 个工作日内交付'],
        ['合同期限', '自生效日起 3 个月，到期前 15 日双方协商续约'],
    ],
    col_widths=[3, 11]
)
add_line()

# 四、双方权责
add_heading_styled('四、双方权责', level=2)

add_heading_styled('甲方责任', level=3)
for i, item in enumerate([
    '按约定时间完成业务对接与内容交付',
    '抖音视频日更 + 投流计划优化（计划维护除外）',
    '达人探店内容发布前与乙方确认脚本与排期',
    '来客后台日常维护，包括团购套餐更新、评价回复、数据报表',
    '对乙方经营数据严格保密',
], 1):
    add_para(f'{i}. {item}')

add_heading_styled('乙方责任', level=3)
for i, item in enumerate([
    '按时支付服务费用',
    '提供必要的门店访问权限（探店拍摄、新品拍摄）',
    '配合提供菜品信息、门店素材、活动信息',
    '来客后台商家账号授权给甲方操作',
    '平台投流广告费由乙方自行充值到广告户，甲方负责投放操作与优化',
], 1):
    add_para(f'{i}. {item}')
add_line()

# 五、成果标准
add_heading_styled('五、成果标准', level=2)
add_table(
    ['服务项', '交付标准'],
    [
        ['达人探店', '每组探店完成拍摄 + 发布'],
        ['矩阵种草', '18 个视频，覆盖多个账号分发'],
        ['视频素材', '100–200 条成品素材，1080P 以上画质'],
        ['来客后台', '工作日日常维护，48 小时内处理差评'],
        ['新品拍摄', '新品上架后 5 个工作日内完成拍摄'],
    ],
    col_widths=[3, 11]
)
add_line()

# 六、合同终止与退款
add_heading_styled('六、合同终止与退款', level=2)
for i, item in enumerate([
    '乙方单方解除：已履约部分不退款，未执行的达人探店 / 矩阵种草按比例折抵。',
    '甲方未按约定标准交付的，乙方有权要求补做或按未交付比例退款。',
    '团购搭建服务费为一次性服务，完成后不予退款。',
    '双方协商一致可提前终止，按实际已服务天数结算。',
], 1):
    add_para(f'{i}. {item}')
add_line()

# 七、其他
add_heading_styled('七、其他', level=2)
for item in [
    '本合同一式两份，双方各执一份，具有同等法律效力。',
    '执行过程中如有争议，双方友好协商解决；协商不成的，提交甲方所在地法院管辖。',
    '合同附件（服务排期表、达人名单）与本合同具有同等效力。',
]:
    p = doc.add_paragraph()
    run = p.add_run('• ' + item)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
add_line()

# 八、签字页
add_heading_styled('八、签字页', level=2)

table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

sign_data = [
    ['甲方（盖章）：昕冠云数', '乙方（盖章/签字）：陶香居私房菜'],
    ['代表签字：_______________', '签字：_______________'],
    ['日期：2026年____月____日', '联系电话：_______________'],
    ['', '日期：2026年____月____日'],
]

for r, row_data in enumerate(sign_data):
    for c, val in enumerate(row_data):
        set_cell_text(table.rows[r].cells[c], val, bold=(r == 0), font_size=11,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER if r == 0 else None)

# 保存
output_path = r'E:\hermes\销冠数字员工\客户合同\陶香居私房菜-服务合同.docx'
doc.save(output_path)
print(f'已生成: {output_path}')
