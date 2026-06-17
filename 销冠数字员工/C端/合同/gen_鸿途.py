"""鸿途新能源合同 Word 生成 - 抖音获客系统"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def _set_run_font(run, size=10.5, bold=False):
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold

def A(doc, text, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run(text), size=size, bold=True)

def B(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    _set_run_font(p.add_run(text), size=size, bold=bold)

def S(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p.add_run('——' * 24))

def H(doc, text):
    p = doc.add_paragraph()
    _set_run_font(p.add_run(text), size=14, bold=True)

def _C(cell, text, bold=False, size=9):
    cell.text = ''
    _set_run_font(cell.paragraphs[0].add_run(text), size=size, bold=bold)

def _shade(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def _border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr or parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def T(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _C(cell, h, bold=True)
        _shade(cell, 'D9E2F3')
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            _C(table.rows[ri + 1].cells[ci], val)
    if col_widths:
        for ri in range(len(rows) + 1):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = Cm(w)
    _border(table)
    doc.add_paragraph()

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

A(doc, '鸿途新能源 · 服务合同')
S(doc)
B(doc, '合同编号：XG-2026-0614', bold=True)
B(doc, '签署日期：2026年____月____日', bold=True)
S(doc)
B(doc, '甲方（服务方）：昕冠科技', bold=True)
B(doc, '乙方（客户方）：鸿途新能源', bold=True)
S(doc)

H(doc, '一、服务内容')
B(doc, '甲方为乙方提供抖音获客系统服务：')

T(doc,
    ['服务模块', '说明'],
    [
        ['视频获客', '数字人口播视频自动生产+多平台一键发布'],
        ['广告投流', '广告计划自动搭建+盯盘实时优化'],
        ['SEO搜索', '百度/地图/抖音等搜索引擎排名优化'],
        ['GEO推荐', '豆包/DeepSeek/Kimi等搜索推荐覆盖'],
        ['来客后台维护', '抖音来客商家后台日常运营维护'],
    ],
    col_widths=[3, 12]
)
S(doc)

H(doc, '二、费用与支付')
B(doc, '费用明细', bold=True)
T(doc,
    ['费用项', '金额', '说明'],
    [
        ['月度服务费', '6000元/月', '含上述全部抖音获客系统服务，无额外费用'],
    ],
    col_widths=[4, 4, 7]
)

B(doc, '支付方式', bold=True)
B(doc, '签约周期：月签')
B(doc, '首期应付：6000元')
B(doc, '续费：到期前7日支付下一周期费用')
B(doc, '支付账户：_______________')
S(doc)

H(doc, '三、服务周期')
B(doc, '合同生效：首笔款项到账之日起')
B(doc, '系统部署：合同生效后7个工作日内完成')
B(doc, '合同期限：自生效日起1个月。到期前15日双方协商续约。')
S(doc)

H(doc, '四、双方权责')
B(doc, '甲方责任：', bold=True)
for item in [
    '按约定时间完成系统部署',
    '保证系统7×24小时正常运行（计划维护除外）',
    '来客后台日常维护，包括团购套餐更新、评价回复、数据报表',
    '对乙方经营数据严格保密',
]:
    B(doc, item)

B(doc, '')
B(doc, '乙方责任：', bold=True)
for item in [
    '按时支付服务费用',
    '配合提供企业信息、产品素材',
    '来客后台商家账号授权给甲方操作',
    '平台投流广告费由乙方自行充值到广告户，甲方负责投放操作与优化',
]:
    B(doc, item)
S(doc)

H(doc, '五、合同终止与退款')
for item in [
    '乙方单方解除：已履约部分不退款。',
    '甲方未按约定标准交付的，乙方有权要求补做或按未交付比例退款。',
    '双方协商一致可提前终止，按实际已服务天数结算。',
]:
    B(doc, item)
S(doc)

H(doc, '六、其他')
for item in [
    '本合同一式两份，双方各执一份，具有同等法律效力。',
    '执行过程中如有争议，双方友好协商解决；协商不成的，提交甲方所在地法院管辖。',
]:
    B(doc, item)
S(doc)

A(doc, '签字页', size=16)
B(doc, '')
B(doc, '甲方（盖章）：昕冠科技', bold=True)
B(doc, '代表签字：_______________')
B(doc, '日期：2026年____月____日')
B(doc, '')
B(doc, '乙方（盖章/签字）：鸿途新能源', bold=True)
B(doc, '签字：_______________')
B(doc, '联系电话：_______________')
B(doc, '日期：2026年____月____日')

output = r'E:\hermes\销冠数字员工\C端\合同\鸿途新能源-服务合同.docx'
doc.save(output)
print(f'已生成: {output}')
