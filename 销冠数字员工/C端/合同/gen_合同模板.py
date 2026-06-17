"""通用合同 Word 模板生成器

用法：
  1. 复制本文件，修改下方 CLIENT_DATA 字典
  2. 运行：python gen_contract.py
  3. 输出 .docx 到同目录
"""

# ═══ 修改这里 ═══
CLIENT_DATA = {
    "name": "【客户名称】",       # 客户公司名
    "contract_no": "XG-2026-____",  # 合同编号
    "price": "【金额】",           # 月费数字（不含"元"）
    "period": "【月签/季度签/半年签/年签】",  # 签约周期
    "months": "【N】",             # 合同月数
    "first_payment": "【金额】",   # 首期应付
    "has_group_buy": False,        # 是否含团购搭建服务费（仅餐饮客户）
    "extras": [],                  # 额外服务说明（如"达人探店 X组"）
    "output": "合同-模板.docx",
}
# ══════════════════

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def _F(run, size=10.5, bold=False):
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold

def title(doc, text, sz=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _F(p.add_run(text), size=sz, bold=True)

def body(doc, text, bold=False):
    p = doc.add_paragraph()
    _F(p.add_run(text), bold=bold)

def sep(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _F(p.add_run('——' * 24))

def heading(doc, text):
    p = doc.add_paragraph()
    _F(p.add_run(text), size=14, bold=True)

def _cell(cell, text, bold=False):
    cell.text = ''
    _F(cell.paragraphs[0].add_run(text), size=9, bold=bold)

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        _cell(c, h, bold=True)
        s = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        c._tc.get_or_add_tcPr().append(s)
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            _cell(t.rows[ri + 1].cells[ci], v)
    if widths:
        for ri in range(len(rows) + 1):
            for ci, w in enumerate(widths):
                t.rows[ri].cells[ci].width = Cm(w)
    b = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    t._tbl.tblPr.append(b)
    doc.add_paragraph()

def generate(data):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)
    st = doc.styles['Normal']
    st.font.name = '宋体'
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题
    title(doc, f'{data["name"]} · 服务合同')
    sep(doc)
    body(doc, f'合同编号：{data["contract_no"]}', bold=True)
    body(doc, '签署日期：2026年____月____日', bold=True)
    sep(doc)
    body(doc, '甲方（服务方）：昕冠科技', bold=True)
    body(doc, f'乙方（客户方）：{data["name"]}', bold=True)
    sep(doc)

    # 一、服务内容
    heading(doc, '一、服务内容')
    body(doc, '甲方为乙方提供抖音获客系统服务：')
    table(doc,
        ['服务模块', '说明'],
        [
            ['视频获客', '数字人口播视频自动生产+多平台一键发布'],
            ['广告投流', '广告计划自动搭建+盯盘实时优化'],
            ['SEO搜索', '百度/地图/抖音等搜索引擎排名优化'],
            ['GEO推荐', '豆包/DeepSeek/Kimi等搜索推荐覆盖'],
            ['来客后台维护', '抖音来客商家后台日常运营维护'],
        ],
        widths=[3, 12]
    )
    for extra in data.get('extras', []):
        body(doc, extra)
    sep(doc)

    # 二、费用
    heading(doc, '二、费用与支付')
    body(doc, '费用明细', bold=True)
    fee_rows = [
        ['月度服务费', f'{data["price"]}元/月', '含上述全部抖音获客系统服务，无额外费用'],
    ]
    if data.get('has_group_buy'):
        fee_rows.append(['团购搭建服务费', '3000元', '抖音/美团等平台团购套餐搭建+上架（一次性）'])
    table(doc, ['费用项', '金额', '说明'], fee_rows, widths=[4, 4, 7])

    body(doc, '支付方式', bold=True)
    body(doc, f'签约周期：{data["period"]}')
    body(doc, f'首期应付：{data["first_payment"]}元')
    body(doc, '续费：到期前7日支付下一周期费用')
    body(doc, '支付账户：_______________')
    sep(doc)

    # 三、服务周期
    heading(doc, '三、服务周期')
    body(doc, '合同生效：首笔款项到账之日起')
    body(doc, '系统部署：合同生效后7个工作日内完成')
    body(doc, f'合同期限：自生效日起{data["months"]}个月。到期前15日双方协商续约。')
    sep(doc)

    # 四、双方权责
    heading(doc, '四、双方权责')
    body(doc, '甲方责任：', bold=True)
    for item in [
        '按约定时间完成系统部署',
        '保证系统7×24小时正常运行（计划维护除外）',
        '来客后台日常维护，包括团购套餐更新、评价回复、数据报表',
        '对乙方经营数据严格保密',
    ]:
        body(doc, item)
    body(doc, '')
    body(doc, '乙方责任：', bold=True)
    for item in [
        '按时支付服务费用',
        '配合提供企业信息、产品素材',
        '来客后台商家账号授权给甲方操作',
        '平台投流广告费由乙方自行充值到广告户，甲方负责投放操作与优化',
    ]:
        body(doc, item)
    sep(doc)

    # 五、合同终止
    heading(doc, '五、合同终止与退款')
    for item in [
        '乙方单方解除：已履约部分不退款。',
        '甲方未按约定标准交付的，乙方有权要求补做或按未交付比例退款。',
        '双方协商一致可提前终止，按实际已服务天数结算。',
    ]:
        body(doc, item)
    sep(doc)

    # 六、其他
    heading(doc, '六、其他')
    for item in [
        '本合同一式两份，双方各执一份，具有同等法律效力。',
        '执行过程中如有争议，双方友好协商解决；协商不成的，提交甲方所在地法院管辖。',
    ]:
        body(doc, item)
    sep(doc)

    # 签字页
    title(doc, '签字页', sz=16)
    body(doc, '')
    body(doc, '甲方（盖章）：昕冠科技', bold=True)
    body(doc, '代表签字：_______________')
    body(doc, '日期：2026年____月____日')
    body(doc, '')
    body(doc, f'乙方（盖章/签字）：{data["name"]}', bold=True)
    body(doc, '签字：_______________')
    body(doc, '联系电话：_______________')
    body(doc, '日期：2026年____月____日')

    doc.save(data['output'])
    print(f'已生成: {data["output"]}')

if __name__ == '__main__':
    generate(CLIENT_DATA)
